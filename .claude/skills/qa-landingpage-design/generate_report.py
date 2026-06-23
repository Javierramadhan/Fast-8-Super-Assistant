"""
qa-design report generator
Modes:
  --screenshot-only  --url URL --session NAME        Take production screenshots (desktop + mobile)
  --save-figma       --node-id ID --output FILE.png  Save a Figma section screenshot via Desktop MCP
  --get-figma-specs  --node-id ID --output FILE.txt  Get design specs for a Figma section
  --report           --input issues.json --output FILE.docx  Generate Word QA report
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from io import BytesIO

# ── colour palette ────────────────────────────────────────────────────────────
RED   = (220, 38, 38)    # production annotation box
GREEN = (22, 163, 74)    # figma annotation box
BOX_WIDTH = 4
FIGMA_MCP_URL = "http://127.0.0.1:3845/mcp"


# ─────────────────────────────────────────────────────────────────────────────
# FIGMA DESKTOP MCP SESSION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _mcp_session():
    """Initialize a Figma Desktop MCP session. Returns (requests.Session, session_id)."""
    try:
        import requests as req
    except ImportError:
        print("ERROR: requests not installed. Run: pip install requests")
        sys.exit(1)

    s = req.Session()
    r = s.post(FIGMA_MCP_URL, json={
        "jsonrpc": "2.0", "id": 1, "method": "initialize",
        "params": {
            "protocolVersion": "2024-11-05",
            "capabilities": {},
            "clientInfo": {"name": "qa-tool", "version": "2.0"}
        }
    }, headers={"Accept": "application/json, text/event-stream"}, timeout=10)

    if r.status_code != 200:
        print(f"ERROR: Figma Desktop MCP returned {r.status_code}. Is Figma Desktop open?")
        sys.exit(1)

    sid = r.headers.get("mcp-session-id") or r.headers.get("Mcp-Session-Id")
    if not sid:
        print("ERROR: No session ID from Figma Desktop MCP. Make sure Figma Desktop is running.")
        sys.exit(1)

    hdrs = {"Accept": "application/json, text/event-stream", "Mcp-Session-Id": sid}
    s.post(FIGMA_MCP_URL, json={"jsonrpc": "2.0", "method": "notifications/initialized"},
           headers=hdrs, timeout=5)
    return s, sid, hdrs


def _mcp_call(session, hdrs, req_id: int, tool: str, args: dict) -> dict:
    """Call a Figma Desktop MCP tool. Returns the parsed result dict."""
    r = session.post(FIGMA_MCP_URL, json={
        "jsonrpc": "2.0", "id": req_id, "method": "tools/call",
        "params": {"name": tool, "arguments": args}
    }, headers=hdrs, timeout=30)

    for line in r.text.split("\n"):
        if line.startswith("data:"):
            try:
                return json.loads(line[5:].strip())
            except Exception:
                pass
    return {}


# ─────────────────────────────────────────────────────────────────────────────
# MODE: SAVE FIGMA SECTION SCREENSHOT
# ─────────────────────────────────────────────────────────────────────────────

def save_figma_section(node_id: str, output_path: str):
    """
    Save a high-resolution screenshot of a specific Figma node (section/frame)
    via the Figma Desktop MCP. Requires Figma Desktop to be open with the file active.
    """
    import base64

    print(f"Connecting to Figma Desktop MCP for node: {node_id}")
    s, sid, hdrs = _mcp_session()

    result = _mcp_call(s, hdrs, req_id=2, tool="get_screenshot", args={"nodeId": node_id})

    content = result.get("result", {}).get("content", [])
    for item in content:
        if item.get("type") == "image":
            img_bytes = base64.b64decode(item["data"])
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            Path(output_path).write_bytes(img_bytes)
            from PIL import Image
            w, h = Image.open(BytesIO(img_bytes)).size
            print(f"Saved: {output_path} ({w}x{h}px, {len(img_bytes)//1024}KB)")
            return
        elif item.get("type") == "text":
            print(f"ERROR from Figma: {item['text']}")
            print("Make sure the Figma file is open and this node ID is in the active tab.")
            sys.exit(1)

    print("No image data received from Figma Desktop MCP.")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────────────────────
# MODE: GET FIGMA DESIGN SPECS (per-section, avoids large-response drop)
# ─────────────────────────────────────────────────────────────────────────────

def get_figma_specs(node_id: str, output_path: str):
    """
    Extract design specs (padding, font, colors) for a specific Figma node
    via the Figma Desktop MCP. Call on SECTION nodes, not the full page frame.
    """
    print(f"Getting design specs for node: {node_id}")
    s, sid, hdrs = _mcp_session()

    result = _mcp_call(s, hdrs, req_id=2, tool="get_design_context", args={
        "nodeId": node_id,
        "artifactType": "WEB_PAGE_OR_APP_SCREEN"
    })

    content = result.get("result", {}).get("content", [])
    specs_text = ""
    for item in content:
        if item.get("type") == "text":
            specs_text += item["text"] + "\n"

    if not specs_text:
        print("No design specs received. Try a smaller node (section/component, not full page).")
        sys.exit(1)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(specs_text, encoding="utf-8")
    print(f"Saved specs: {output_path} ({len(specs_text)} chars)")
    print("\n--- PREVIEW (first 800 chars) ---")
    print(specs_text[:800])


# ─────────────────────────────────────────────────────────────────────────────
# MODE: TAKE PRODUCTION SCREENSHOTS
# ─────────────────────────────────────────────────────────────────────────────

def take_screenshots(url: str, session: str) -> dict:
    """
    Capture desktop (1440px) and mobile (375px) screenshots via Playwright.
    Uses device_scale_factor=1 to guarantee correct pixel dimensions on HiDPI displays.
    Uses wait_until='load' (not 'networkidle') to handle SPAs that never fully idle.
    """
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("ERROR: Playwright not installed. Run: pip install playwright && python -m playwright install chromium")
        sys.exit(1)

    skill_dir = Path(__file__).parent
    out_dir = skill_dir / "reports" / "screenshots" / session
    out_dir.mkdir(parents=True, exist_ok=True)

    viewports = [
        ("desktop_1440", 1440, 900),
        ("mobile_375",   375,  812),
    ]

    paths = {}
    with sync_playwright() as p:
        browser = p.chromium.launch()
        for name, width, height in viewports:
            # device_scale_factor=1 fixes HiDPI pixel doubling (e.g. 558px instead of 375px)
            context = browser.new_context(
                viewport={"width": width, "height": height},
                device_scale_factor=1
            )
            page = context.new_page()
            page.goto(url, wait_until="load", timeout=30000)
            page.wait_for_timeout(2500)   # let lazy images/fonts settle
            out_path = out_dir / f"{name}.png"
            page.screenshot(path=str(out_path), full_page=True)

            from PIL import Image as _Img
            actual_w, actual_h = _Img.open(out_path).size
            print(f"  {name}: {actual_w}x{actual_h}px → {out_path.name}")

            paths[name] = str(out_path)
            page.close()
            context.close()
        browser.close()

    print(f"\nScreenshots saved to: {out_dir}")
    return paths


# ─────────────────────────────────────────────────────────────────────────────
# IMAGE ANNOTATION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _load_img(path: str):
    from PIL import Image
    return Image.open(path).convert("RGB")


def annotate_and_crop(img_path: str, crop: list | None, box_color: tuple) -> bytes:
    """
    Crop the image to `crop` [x, y, w, h], draw a coloured border, return PNG bytes.
    Adds 30px context padding around the crop.
    If crop is None, downsamples the full page to max 1200px tall.
    """
    from PIL import ImageDraw

    img = _load_img(img_path)
    iw, ih = img.size

    if crop:
        x, y, w, h = [int(v) for v in crop]
        x  = max(0, min(x, iw));  y  = max(0, min(y, ih))
        x2 = max(0, min(x+w, iw)); y2 = max(0, min(y+h, ih))
        pad = 40
        cx = max(0, x-pad);  cy = max(0, y-pad)
        cx2 = min(iw, x2+pad); cy2 = min(ih, y2+pad)
        region = img.crop((cx, cy, cx2, cy2))
        draw = ImageDraw.Draw(region)
        draw.rectangle([x-cx, y-cy, x2-cx, y2-cy], outline=box_color, width=BOX_WIDTH)
        result = region
    else:
        max_h = 1400
        if ih > max_h:
            ratio = max_h / ih
            result = img.resize((int(iw*ratio), max_h), resample=3)
        else:
            result = img

    buf = BytesIO()
    result.save(buf, format="PNG")
    buf.seek(0)
    return buf.read()


def _fit_image_to_column(img_bytes: bytes, col_width_inches: float = 3.0) -> tuple:
    """
    Return (width_inches, height_inches) that fits the image in col_width_inches
    while preserving aspect ratio. Caps height at 4 inches to avoid giant images.
    """
    from PIL import Image
    img = Image.open(BytesIO(img_bytes))
    w, h = img.size
    ratio = h / w
    final_w = col_width_inches
    final_h = final_w * ratio
    if final_h > 4.0:   # cap tall screenshots
        final_h = 4.0
        final_w = final_h / ratio
    return final_w, final_h


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def add_header(doc, product: str, url: str, date: str, qa_by: str):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_heading("", level=0)
    title.clear()
    run = title.add_run(f"QA Design Report — {product}")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(15, 23, 42)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run("URL: ").bold = True;       meta.add_run(url)
    meta.add_run("   |   Date: ").bold = True; meta.add_run(date)
    meta.add_run("   |   QA by: ").bold = True; meta.add_run(qa_by)
    meta.add_run("   |   Generated: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    doc.add_paragraph()


def add_section_heading(doc, text: str):
    from docx.shared import Pt, RGBColor
    h = doc.add_heading("", level=1)
    h.clear()
    run = h.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 64, 175)
    doc.add_paragraph()


def _severity_color(s: str) -> tuple:
    return {"critical": (220,38,38), "major": (234,88,12), "minor": (100,116,139)}.get(s.lower(), (100,116,139))


def _set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_issue_block(doc, issue: dict, issue_dir: Path):
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sev   = issue.get("severity", "minor").upper()
    cat   = issue.get("category", "design").upper()
    idx   = issue.get("id", "?")
    title = issue.get("title", "Untitled Issue")

    # ── Issue header ──
    hdr = doc.add_paragraph()
    r = hdr.add_run(f"#{idx}  ");  r.bold = True; r.font.size = Pt(11)
    r = hdr.add_run(title + "   "); r.bold = True; r.font.size = Pt(11)
    sc = _severity_color(sev)
    r = hdr.add_run(f"[{sev}]  ");  r.bold = True;  r.font.color.rgb = RGBColor(*sc)
    r = hdr.add_run(f"[{cat}]");    r.font.color.rgb = RGBColor(59, 130, 246)

    loc_p = doc.add_paragraph()
    loc_p.add_run("Viewport: ").bold = True;  loc_p.add_run(issue.get("viewport","desktop") + "   ")
    loc_p.add_run("Location: ").bold = True;  loc_p.add_run(issue.get("location","—"))

    # ── 2-column screenshot table ──
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    _set_cell_bg(table.cell(0,0), "DBEAFE")  # blue header — production
    _set_cell_bg(table.cell(0,1), "DCFCE7")  # green header — figma

    for col_idx, label in enumerate(["PRODUCTION", "FIGMA"]):
        p = table.cell(0, col_idx).paragraphs[0]; p.clear()
        r = p.add_run(label); r.bold = True; r.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Production screenshot
    cell_prod = table.cell(1, 0)
    prod_ss   = issue.get("production_screenshot")
    prod_crop = issue.get("production_crop")
    if prod_ss and os.path.exists(prod_ss):
        try:
            img_bytes = annotate_and_crop(prod_ss, prod_crop, RED)
            w_in, h_in = _fit_image_to_column(img_bytes, 3.0)
            p = cell_prod.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(BytesIO(img_bytes), width=Inches(w_in))
        except Exception as e:
            cell_prod.paragraphs[0].add_run(f"[Error: {e}]")
    else:
        cell_prod.paragraphs[0].add_run("[No production screenshot]")

    # Figma screenshot
    cell_fig = table.cell(1, 1)
    fig_ss   = issue.get("figma_screenshot")
    fig_crop = issue.get("figma_crop")
    if fig_ss and os.path.exists(fig_ss):
        try:
            img_bytes = annotate_and_crop(fig_ss, fig_crop, GREEN)
            w_in, h_in = _fit_image_to_column(img_bytes, 3.0)
            p = cell_fig.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(BytesIO(img_bytes), width=Inches(w_in))
        except Exception as e:
            cell_fig.paragraphs[0].add_run(f"[Error: {e}]")
    else:
        cell_fig.paragraphs[0].add_run("[No Figma screenshot — lihat deskripsi]")

    # ── Description ──
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run("Masalah: ").bold = True;    p.add_run(issue.get("problem","—"))
    p = doc.add_paragraph(); p.add_run("Seharusnya: ").bold = True; p.add_run(issue.get("should_be","—"))

    # Separator
    sep = doc.add_paragraph("─" * 80)
    sep.runs[0].font.color.rgb = RGBColor(203,213,225)
    sep.runs[0].font.size = Pt(7)
    doc.add_paragraph()


def add_summary_table(doc, issues: list):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    h = doc.add_heading("", level=1); h.clear()
    r = h.add_run("Summary"); r.font.size = Pt(14); r.font.color.rgb = RGBColor(30,64,175)

    cats = ["content","design"]
    sevs = ["critical","major","minor"]
    counts = {c: {s: 0 for s in sevs} for c in cats}
    for issue in issues:
        c = issue.get("category","design").lower()
        s = issue.get("severity","minor").lower()
        if c in counts and s in sevs:
            counts[c][s] += 1

    table = doc.add_table(rows=len(cats)+2, cols=5)
    table.style = "Table Grid"

    # Header row with background
    _set_cell_bg(table.rows[0].cells[0], "1E3A8A")  # dark blue header
    for col, text in enumerate(["Category","Critical","Major","Minor","Total"]):
        p = table.cell(0,col).paragraphs[0]; p.clear()
        r = p.add_run(text); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255,255,255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col > 0:
            _set_cell_bg(table.cell(0,col), "1E3A8A")

    totals = {s: 0 for s in sevs}
    for row_i, cat in enumerate(cats, start=1):
        row_total = sum(counts[cat].values())
        bg = "EFF6FF" if row_i % 2 == 0 else "FFFFFF"
        vals = [cat.capitalize()] + [str(counts[cat][s]) for s in sevs] + [str(row_total)]
        for col, val in enumerate(vals):
            _set_cell_bg(table.cell(row_i, col), bg)
            p = table.cell(row_i,col).paragraphs[0]; p.clear()
            p.add_run(val).font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for s in sevs:
            totals[s] += counts[cat][s]

    grand = sum(totals.values())
    last_row = len(cats)+1
    vals = ["TOTAL"] + [str(totals[s]) for s in sevs] + [str(grand)]
    for col, val in enumerate(vals):
        _set_cell_bg(table.cell(last_row, col), "DBEAFE")
        p = table.cell(last_row, col).paragraphs[0]; p.clear()
        r = p.add_run(val); r.bold = True; r.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─────────────────────────────────────────────────────────────────────────────
# MODE: GENERATE WORD REPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(issues_path: str, output_path: str):
    from docx import Document
    from docx.shared import Cm

    with open(issues_path, encoding="utf-8") as f:
        data = json.load(f)

    issues = data.get("issues", [])
    if not issues:
        print("No issues found in issues.json.")
        return

    issue_dir = Path(issues_path).parent
    for issue in issues:
        for key in ("production_screenshot", "figma_screenshot"):
            val = issue.get(key)
            if val and not os.path.isabs(val):
                issue[key] = str(issue_dir / val)

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Cm(1.8)
        section.top_margin = section.bottom_margin = Cm(1.8)

    add_header(doc,
        product=data.get("product","Product"),
        url=data.get("url","—"),
        date=data.get("date", datetime.now().strftime("%Y-%m-%d")),
        qa_by=data.get("qa_by","Javier Ramadhan"))

    content_issues = [i for i in issues if i.get("category","").lower() == "content"]
    design_issues  = [i for i in issues if i.get("category","").lower() == "design"]

    if content_issues:
        add_section_heading(doc, "1. Content / Writing Issues")
        for issue in content_issues:
            add_issue_block(doc, issue, issue_dir)

    if design_issues:
        add_section_heading(doc, "2. Design Properties (Padding & Responsive)")
        for issue in design_issues:
            add_issue_block(doc, issue, issue_dir)

    add_summary_table(doc, issues)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"\nReport saved: {out}")
    print(f"Total: {len(issues)} issues ({len(content_issues)} content, {len(design_issues)} design)")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="QA Design Report Generator v2")
    parser.add_argument("--screenshot-only", action="store_true", dest="screenshot_only")
    parser.add_argument("--save-figma",      action="store_true", dest="save_figma")
    parser.add_argument("--get-figma-specs", action="store_true", dest="get_figma_specs")
    parser.add_argument("--report",          action="store_true", dest="report")

    parser.add_argument("--url",     type=str)
    parser.add_argument("--session", type=str)
    parser.add_argument("--node-id", type=str, dest="node_id")
    parser.add_argument("--output",  type=str)
    parser.add_argument("--input",   type=str)

    args = parser.parse_args()

    if args.screenshot_only:
        if not args.url or not args.session:
            print("Usage: --screenshot-only --url URL --session SESSION_NAME")
            sys.exit(1)
        take_screenshots(args.url, args.session)

    elif args.save_figma:
        if not args.node_id or not args.output:
            print("Usage: --save-figma --node-id NODE_ID --output path/to/section.png")
            sys.exit(1)
        save_figma_section(args.node_id, args.output)

    elif args.get_figma_specs:
        if not args.node_id or not args.output:
            print("Usage: --get-figma-specs --node-id NODE_ID --output path/to/specs.txt")
            sys.exit(1)
        get_figma_specs(args.node_id, args.output)

    elif args.report:
        if not args.input or not args.output:
            print("Usage: --report --input issues.json --output report.docx")
            sys.exit(1)
        generate_report(args.input, args.output)

    else:
        parser.print_help()


if __name__ == "__main__":
    main()
